import os
import sys
from importlib.resources import as_file, files

import docx
import numpy as np
import pytest
import pyvista as pv

import tests.resources.plotter as pkg_res
from f4enix.constants import ITER_Z_LEVELS
from f4enix.output.plotter import Atlas, CDFplot, MeshPlotter

RESOURCES = files(pkg_res)

# Do not know how to implement the root test without having a
# real path
RESOURCES_PATH = os.path.dirname(os.path.abspath(pkg_res.__file__))


@pytest.fixture
def plotter() -> MeshPlotter:
    with as_file(RESOURCES.joinpath("meshtal_1004_vtk.vtr")) as infile:
        mesh = pv.read(infile)
    with as_file(RESOURCES.joinpath("iter1D.stl")) as infile:
        stl = pv.read(infile)

    plotter = MeshPlotter(mesh, stl.scale(10))

    return plotter


@pytest.fixture
def plotter_no_stl() -> MeshPlotter:
    with as_file(RESOURCES.joinpath("meshtal_1004_vtk.vtr")) as infile:
        mesh = pv.read(infile)

    plotter = MeshPlotter(mesh)

    return plotter


class TestMeshPlotter:
    def test_slice_toroidal(self, plotter: MeshPlotter, plotter_no_stl: MeshPlotter):
        slices = plotter.slice_toroidal(20)
        assert len(slices) == 9

        # Check that they are not empty
        for slice in slices:
            assert slice[1].bounds is not None
            assert slice[2].bounds is not None

        slices = plotter_no_stl.slice_toroidal(20)
        assert len(slices) == 9

        # Check that they are not empty
        for slice in slices:
            assert slice[1].bounds is not None

    @pytest.mark.parametrize("axis", ["x", "y", "z"])
    def test_slice_on_axis(
        self, plotter: MeshPlotter, plotter_no_stl: MeshPlotter, axis
    ):
        slices = plotter.slice_on_axis(axis, 3)
        # Check that they are not empty
        for slice in slices:
            assert slice[1].bounds is not None
            assert slice[2].bounds is not None

        slices = plotter_no_stl.slice_on_axis(axis, 3)
        # Check that they are not empty
        for slice in slices:
            assert slice[1].bounds is not None

    @pytest.mark.skipif(sys.platform == "win32", reason="Windows access error")
    def test_plot_slices(
        self, plotter: MeshPlotter, plotter_no_stl: MeshPlotter, tmpdir
    ):
        slices = plotter.slice_on_axis("y", 5)[1:-1]
        outpath = tmpdir.mkdir("meshplotter")
        plotter.plot_slices(slices, "Value - Total", outpath=outpath)
        assert len(os.listdir(outpath)) == 3

        slices = plotter_no_stl.slice_on_axis("y", 5)[1:-1]
        im = plotter_no_stl.plot_slices(slices, "Value - Total")
        assert len(im) == 3

        # test with categories
        slices = plotter_no_stl.slice_on_axis("y", 5)[1:-1]
        im = plotter_no_stl.plot_slices(
            slices, "Value - Total", custom_categories="TNF"
        )
        assert len(im) == 3

    def test_slice(self, plotter: MeshPlotter, plotter_no_stl: MeshPlotter):
        slices = plotter.slice(
            [["slice 1", 0, 0, 0, -0.1, 0.5, 1], ["slice 2", 0, 15, 0, 0.1, 0.5, 1]]
        )
        # verify that they intersect
        for slice in slices:
            assert slice[-1] is not None
            assert slice[1].bounds is not None

        # coordinates are in meters, input in mm
        slices = plotter_no_stl.slice(
            [["slice 1", 0, 0, 0, -0.1, 0.5, 1], ["slice 2", 0, 15, 0, 0.1, 0.5, 1]]
        )
        # verify that they intersect
        for slice in slices:
            assert slice[1].bounds is not None


class TestAtlas:
    def test_build_from_root(self, tmpdir):
        # cannot run this test on linux architecture
        name = "test"
        atlas = Atlas(name=name)
        atlas.build_from_root(os.path.join(RESOURCES_PATH, "root"))
        outfolder = tmpdir.mkdir("atlas")
        try:
            atlas.save(outfolder)

            # try to open it
            with open(os.path.join(outfolder, name + ".docx"), "rb") as infile:
                doc = docx.Document(infile)
                assert len(doc.paragraphs) == 15

        except NotImplementedError:
            # cannot be tested if word is not installed
            assert True

    @pytest.mark.skipif(sys.platform == "win32", reason="Windows access error")
    def test_add_section(self, tmpdir, plotter):
        name = "test2"
        atlas = Atlas(name=name)

        # Get the slices
        slices = plotter.slice_on_axis("x", 3)
        images = plotter.plot_slices(slices, "Value - Total")
        atlas.add_section("New section", images)

        outfolder = tmpdir.mkdir("atlas2")

        assert len(atlas.doc.paragraphs) == 9

        atlas.save(outfolder)


class TestCDFplot:
    def test_plot(self, tmpdir):
        data_list = []
        labels = []
        for i in range(10):
            data_list.append(np.random.random(100))
            labels.append(str(i))

        plotter = CDFplot(suptitle="title", xlabel="xlabel", ylabel="ylabel")
        plotter.plot(data_list, datalabels=labels)
        plotter.save(os.path.join(tmpdir, "trial.png"))
