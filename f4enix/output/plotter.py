from __future__ import annotations

"""
Plotter classes and methods for neutronics outputs.

For the moment it allows only to plot slices of meshes.
"""

"""
Copyright 2019 F4E | European Joint Undertaking for ITER and the Development of
Fusion Energy (‘Fusion for Energy’). Licensed under the EUPL, Version 1.2 or - 
as soon they will be approved by the European Commission - subsequent versions
of the EUPL (the “Licence”). You may not use this work except in compliance
with the Licence. You may obtain a copy of the Licence at: 
    https://eupl.eu/1.2/en/  
Unless required by applicable law or agreed to in writing, software distributed
under the Licence is distributed on an “AS IS” basis, WITHOUT WARRANTIES OR
CONDITIONS OF ANY KIND, either express or implied. See the Licence permissions
and limitations under the Licence.
"""

import os
import pyvista as pv
import numpy as np
import logging
import docx

from typing import Union
from docx.shared import Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from PIL import Image, ImageOps
import io
from copy import deepcopy
from abc import ABC, abstractmethod
import matplotlib.pyplot as plt
from matplotlib.ticker import PercentFormatter

from math import radians, degrees
from f4enix.constants import TID_CATEGORIES, TNF_CATEGORIES, SDDR_CATEGORIES

pv.set_plot_theme("document")


class MeshPlotter:

    def __init__(self, mesh: pv.PolyData, stl: pv.PolyData = None) -> None:
        """Object responsible for the plotting of meshes.

        Allows slicing in different orientation and plotting of slices.

        Parameters
        ----------
        mesh : pv.PolyData
            mesh to be sliced
        stl : pv.PolyData, optional
            stl to be superimposed in the slices, by default None

        Attributes
        ----------
        mesh : pv.PolyData
            mesh to be sliced
        stl : pv.PolyData, optional
            stl to be superimposed in the slices
        legend_args : dict
            contains the default parameters for the plot of the scalar bars.
            use print(self.legend_args) to check the default values. The entire
            list of parameters that can be used can be found in the
            `pv.Plotter.add_scalar_bar() method <https://docs.pyvista.org/version/stable/api/plotting/_autosummary/pyvista.Plotter.add_scalar_bar.html>`_

        Examples
        --------
        Load a mesh and stl file and slice them in different way

        >>> from f4enix.output.plotter import MeshPlotter
        ... # Initialize the plotter with the mesh and stl
        ... # scaling the mesh and stl accordingly to be in the same units
        ... plotter = MeshPlotter(global_mesh.scale(0.01, inplace=False),
        ...                       stl=stl.scale(0.001, inplace=False))
        ... # There are many default settings that can be modified
        ... plotter.legend_args['vertical'] = False
        ... # -- Different slicing methods can be used: --
        ... # toroidal slicing
        ... toroidal_slices = plotter.slice_toroidal(30)  # 30 deg. increment
        ... # vertical slicing on an axis and defininig number of divisions
        ... horizontal_slices = plotter.slice_on_axis('z', 3)
        ... # General slicing using origin and normals for each slice
        ... slice_params = [
                ['slice 1', 0, 0, -11.500, 0, 0, 1],
                ['slice 2', 9, 0, -8.425, 0.25, 0, 0.9]]
        ... general_slices = plotter.slice(slice_params)

        and then plot the meshes

        >>> # Plot the slices
        ... array_name = 'scalar1'  # name of the scalar to be plot from the mesh
        ... images = plotter.plot_slices(horizontal_slices, array_name,
        ...                              n_colors=7,  # colors of the legend
        ...                              min_max=(1e6, 1e12),  # limits
        ...                              scale_title='Gy/h')  # legend title

        """
        self.mesh = mesh

        self.stl = stl

        if stl is not None:
            self._has_stl = True
        else:
            self._has_stl = False

        # Default legend settings
        self.legend_args = dict(
            title_font_size=18,
            label_font_size=18,
            shadow=True,
            # n_labels=7,
            italic=True,
            fmt="%.e",
            font_family="arial",
            vertical=True,
            position_x=0.85,
            position_y=0.25,
        )

    def _get_plotter(self) -> pv.Plotter:
        # Initiate the plotter with all default actions if needed
        pl = pv.Plotter(off_screen=True)
        # TODO horizontal bar may be issue here
        pl.add_axes(viewport=(0.8, 0, 1, 0.2))
        # pl.set_background('white')

        return pl

    def slice(
        self, slice_param: list[list]
    ) -> list[tuple[str, pv.PolyData, Union[pv.PolyData, None]]]:
        """Perform arbitrary general slicing providing origin and normal
        vectors.

        Parameters
        ----------
        slice_param : list[list]
            matrix of parameters to be provided, one row for each slice.
             [[name, x, y, z, normx, normy, normz], ...]

        Returns
        -------
        list[tuple[str, pv.PolyData, Union[pv.PolyData, None]]]
            contains a list of (slice name, mesh slice, stl slice).
            the names are assigned according to slicing logic.
            In case of no stl assigned to the plotter, the stl slice will be
            equal to None.

        """

        outp = []
        for params in slice_param:
            name = params[0]
            origin = params[1:4]
            norm = params[4:]

            mesh_slice = self.mesh.slice(origin=origin, normal=norm)

            # get the corresponding stl slices if needed
            if self._has_stl:
                stl_slice = self._get_stl_slices([mesh_slice])[0]
            else:
                stl_slice = None

            outp.append((name, mesh_slice, stl_slice))

        return outp

    def slice_on_axis(
        self, axis: str, n: int
    ) -> list[tuple[str, pv.PolyData, Union[pv.PolyData, None]]]:
        """Creates a series of slice of the mesh distributed equally along the
        specified axis. Stl file will be sliced accordingly if present

        Parameters
        ----------
        axis : str
            either 'x', 'y' or 'z'
        n : int
            number of divisions

        Returns
        -------
        list[tuple[str, pv.PolyData, Union[pv.PolyData, None]]]
            contains a list of (slice name, mesh slice, stl slice).
            the names are assigned according to slicing logic.
            In case of no stl assigned to the plotter, the stl slice will be
            equal to None.

        """
        # Use the automatic slicing from pyvista
        idxs = {"x": 0, "y": 1, "z": 2}

        # perform the slices
        slices = self.mesh.slice_along_axis(n, axis=axis)
        # get the corresponding stl slices if needed
        if self._has_stl:
            stl_slices = self._get_stl_slices(slices)

        outp = []
        # build the output
        for i, mslice in enumerate(slices):
            # TODO maybe a smarter function here
            name = "P{} = {}".format(axis, round(mslice.center[idxs[axis]], 1))
            if self._has_stl:
                stl_slice = stl_slices[i]
            else:
                stl_slice = None

            outp.append((name, mslice, stl_slice))

        return outp

    def slice_toroidal(
        self,
        theta_increment: float,
        center: list = [0, 0, 0],
        min_max_theta: tuple[float, float] = None,
    ) -> list[tuple[str, pv.PolyData, Union[pv.PolyData, None]]]:
        """Create a series of slices of the mesh distributed toroidally equally
        around the vertical axis (z). Stl file will be sliced accordingly if
        present

        Parameters
        ----------
        theta_increment : float
            theta increments to be performed for each slide
        center : list, optional
            x, y, z point to be used as center for the
            slicing, by default [0, 0, 0].
        min_max_theta : tuple(float, float), optional
            specify a minimum and max theta for the plots (degrees).
            This may help
            avoiding slices in empty regions in partial models if the mesh
            is not properly rotated. By default is None, meaning that the
            slicing will start at theta = 0 and finish at theta = pi.

        Returns
        -------
        list[tuple[str, pv.PolyData, Union[pv.PolyData, None]]]
            contains a list of (slice name, mesh slice, stl slice).
            the names are assigned according to slicing logic.
            In case of no stl assigned to the plotter, the stl slice will be
            equal to None.
        """

        center = np.array(center)
        increment = radians(theta_increment)
        mesh_slices = []

        if min_max_theta is None:
            angles = np.arange(0, np.pi, increment)
        else:
            angles = np.arange(
                radians(min_max_theta[0]), radians(min_max_theta[1]), increment
            )

        for theta in angles:
            normal = np.array([np.cos(theta), np.sin(theta), 0.0]).dot(np.pi / 2.0)

            mesh_slice = self.mesh.slice(origin=center, normal=normal)

            # try to access its normal, if it has it, it is not empty
            try:
                mesh_slice.cell_normals[0]
            except KeyError:
                # it means that nothing can be sliced here because the
                # the slice is empty
                logging.warning(
                    "No slice can be done at theta={} deg".format(degrees(theta))
                )
                continue

            mesh_slices.append(mesh_slice)

        if self._has_stl:
            stl_slices = self._get_stl_slices(mesh_slices)

        outp = []
        # build the output
        for i, mslice in enumerate(mesh_slices):
            name = "theta = {} deg".format(round(degrees(angles[i]), 1))
            if self._has_stl:
                stl_slice = stl_slices[i]
            else:
                stl_slice = None

            outp.append((name, mslice, stl_slice))

        return outp

    def plot_slices(
        self,
        slices: list[tuple[str, pv.PolyData, Union[pv.PolyData, None]]],
        array_name: str,
        outpath: os.PathLike = None,
        min_max: tuple[float] = None,
        log_scale: bool = True,
        stl_color: str = "white",
        n_colors: int = 256,
        scale_quality: float = 3,
        scale_title: str = None,
        custom_categories: str = None,
    ) -> list[tuple[str, Image.Image]]:
        """Plot a series of slices to an outpath folder. The slices names
        are used as file names.

        Parameters
        ----------
        slices : list[tuple[str, pv.PolyData, Union[pv.PolyData, None]]]
            list of slices to be plotted. Usually produced with MeshPlotter
            methods.
        array_name : str
            name of the scalar array to be plotted
        outpath : os.PathLike, optional
            path to the directory that will contain all the plots. This folder
            must be already created. Files with the same name will be
            overridden. default is None, meaning that images will not be saved
        min_max : tuple[float], optional
            min and max values to be set for the scalars in all plots
        log_scale : bool, optional
            if true, activate logarithmic scale, by default True
        stl_color : str, optional
            color for the stl, by default 'white'
        n_colors : int, optional
            number of discrete colors to be used in the legend, by default 256
        scale_quality : float, optional
            increase the resolution of the picture by a factor, by default 2
        scale_title : str, optional
            ovverride the title of the scalar legend. The defualt is None,
            meaning that the array name will be used.
        custom_categorical : str
            plot special categorical plots. Admitted categories are TID,
            TNF and SDDR.

        Returns
        -------
        list[tuple[str, Image.Image]]
            list of images of the slices

        Raises
        ------
        ValueError
            if the path do not exists
        """
        # Check if a categorical plot is requested
        if custom_categories is not None:
            if custom_categories == "SDDR":
                ctg = SDDR_CATEGORIES
            elif custom_categories == "TNF":
                ctg = TNF_CATEGORIES
            elif custom_categories == "TID":
                ctg = TID_CATEGORIES
            else:
                raise ValueError(
                    "{} are not valid custom categories".format(custom_categories)
                )

        # Check that outpath exists
        if outpath is not None:
            if not os.path.exists(outpath):
                raise ValueError("{} does not exists".format(outpath))

        images = []

        scalar_bar_args = deepcopy(self.legend_args)
        if scale_title is not None:
            scalar_bar_args["title"] = scale_title

        # get a number of labels that pairs with the number of colors
        # works only with even numbers
        if n_colors < 30 and n_colors % 2 == 0:
            scalar_bar_args["n_labels"] = int(n_colors / 2 + 1)

        for i, (name, mesh_slice, stl_slice) in enumerate(slices):

            pl = self._get_plotter()
            if custom_categories is not None:
                # Add category label to the mesh
                colors = self._add_categorization(
                    mesh_slice,
                    array_name,
                    ctg["values"],
                    ctg["categories"],
                    ctg["colors"],
                    name=custom_categories,
                )
                scalars = custom_categories
                cmap = colors
                below_color = None
                above_color = None
                lscale = False
                categories = True
            else:
                scalars = array_name
                cmap = "jet"
                below_color = "grey"
                above_color = "purple"
                lscale = log_scale
                categories = False

            pl.add_mesh(
                mesh_slice,
                scalars=scalars,
                scalar_bar_args=scalar_bar_args,
                log_scale=lscale,
                below_color=below_color,
                above_color=above_color,
                clim=min_max,
                cmap=cmap,
                n_colors=n_colors,
                categories=categories,
            )
            if stl_slice is not None:
                pl.add_mesh(stl_slice, color=stl_color)

            # if i == 0:
            #     # ensure that all pictures will have the same bounds
            #     bounds = np.array(mesh_slice.bounds)*0.8
            bounds = np.array(mesh_slice.bounds)

            self._set_perpendicular_camera(mesh_slice, pl, bounds=bounds)

            # Get the image from the plotter and trim it
            im = Image.fromarray(pl.screenshot(None, return_img=True))
            bbox = ImageOps.invert(im).getbbox()
            trimmed = im.crop(bbox)

            images.append((name, trimmed))

            if outpath is not None:
                filename = os.path.join(outpath, "{}.png".format(name))
                pl.screenshot(filename, scale=scale_quality)

        return images

    @staticmethod
    def _set_perpendicular_camera(
        mesh_slice: pv.PolyData, pl: pv.Plotter, bounds: list = None
    ) -> None:
        # align camera: focus on center, position at center + normal
        center = mesh_slice.center
        pl.camera.focal_point = center
        pl.camera.position = center - mesh_slice.cell_normals[0]
        # reset camera to put entire mesh in view
        if bounds is None:
            pl.reset_camera()
        else:
            # help put Y up in case of PZ plots
            if bounds[-2] == bounds[-1]:
                pl.set_viewup([0, 1, 0])
            pl.reset_camera(bounds=bounds)

    def _get_stl_slices(
        self, mesh_slices: list[pv.PolyData]
    ) -> Union[list[pv.PolyData], None]:
        stl_slices = []
        # get the correspondent stl_slices
        for mesh_slice in mesh_slices:
            # check it is not empty
            if len(mesh_slice[mesh_slice.array_names[0]]) == 0:
                stl_slices.append(None)
                continue

            norm = mesh_slice.cell_normals[0]
            stl_slice = self.stl.slice(normal=norm, origin=mesh_slice.center)
            if stl_slice.bounds is None:
                # This may happen if the stl is smaller than the mesh
                stl_slices.append(None)
                continue
            # give a very small translation in order to not be exactly
            # coincident, using the normal should be sufficient
            # but it needs to be properly scaled
            scale = np.abs(mesh_slice.points).mean() * 1e-3
            stl_slice.translate(-norm * scale, inplace=True)
            stl_slices.append(stl_slice)

        return stl_slices

    @staticmethod
    def _add_categorization(
        mesh: pv.PolyData,
        array_name: str,
        vals: list[float],
        categories: list[str],
        colors: list[str],
        name="label",
    ) -> list[str]:

        values = mesh[array_name]
        labels = np.empty(len(values), dtype="<U10")

        # categories = list(range(0, len(vals)+1))
        labels[:] = categories[-1]

        vals.reverse()
        categories.reverse()
        for val, ctg in zip(vals, categories):
            labels[values < val] = ctg

        mesh[name] = labels

        colors_to_use = []
        categories.reverse()
        for category, color in zip(categories, colors):
            if (labels == category).sum() > 1:
                colors_to_use.append(color)

        return colors_to_use

    # @staticmethod
    # def _get_perpendicular_vector(v):
    #     if v[1] == 0 and v[2] == 0:
    #         if v[0] == 0:
    #             raise ValueError('zero vector')
    #         else:
    #             return np.cross(v, [0, 1, 0])

    #     return np.cross(v, [1, 0, 0])


class Atlas:

    def __init__(self, name: str = "atlas", landscape: str = True) -> None:
        """Class to handle the generation of the Atlas.

        Parameters
        ----------
        name : str, optional
            atlas name, by default 'atlas'
        landscape : str, optional
            if true the atlas will be produced in landscape orientation,
            by default True

        Attributes
        ----------
        name : str
            atlas name, by default 'atlas'
        doc : docx.Document
            word document used for the creation of the atlas
        default_width : float
            width used to rescale images that are added to the atlas.
            by default this is set to 0.9*text_length in the document.
            to change it, it is recommended to use either the
            docx.shared.Inches or docx.shared.Inches.Mm conversion method.

        Examples
        --------
        Build an atlas adding sections manually (recommended, since images
        do not need to be saved to the disk). The images can be produced using
        py:method:`f4enix.output.plotter.MeshPlotter.plot_slices`. Both a
        Word and PDF version are saved.

        >>> from f4enix.output.plotter import Atlas
        ... atlas = Atlas('Some title for the atlas')
        ... # reduce for instance the default width for the images
        ... atlas.default_width = atlas.default_width*0.9
        ... # add a section to the atlas containing images produced with
        ... # plot_slices
        ... atlas.add_section('new section', images)
        ... atlas.save('path/to/outfolder')

        """

        self.name = name
        doc = docx.Document()
        doc.add_heading(name, level=0)
        self.doc = doc  # Word Document
        if landscape:
            self._change_orientation()

        # set default width
        section = self.doc.sections[-1]
        margin = 10  # mm
        section.top_margin = Mm(margin)
        section.bottom_margin = Mm(margin)
        section.left_margin = Mm(margin)
        section.right_margin = Mm(margin)
        width = (
            (section.page_width - section.left_margin - section.right_margin)
            / 36000
            * 0.9
        )  # mm
        self.default_width = Mm(width)

    def _insert_img(self, img: os.PathLike, width=None) -> None:
        if width is None:
            width = self.default_width
        self.doc.add_picture(img, width=width)
        # self.doc.add_picture(img, height=Inches(7.5))
        last_paragraph = self.doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_section(
        self,
        section_name: str,
        images: list[tuple[str, Image.Image]],
        level: int = 1,
        include_section_name: bool = True,
        disclaimer: str = None,
    ) -> None:
        """Add a section of plots to the Atlas.

        add a chapter to atlas, the level can be decided.

        Parameters
        ----------
        section_name : str
            Name of the section
        images : list[tuple[str, Image.Image]]
            list of images produced by the plot_slices() method to be added
            to the atlas.
        level : int, optional
            nested level where the section has to be added, by default 1
        include_section_name : bool
            If True, the section name is added to the image name, by default
            True
        disclaimer: str
            add a disclaimer equal for all plots to their titles, by default is
            None

        """

        self.doc.add_heading(section_name, level=level)

        for name, image in images:
            if include_section_name:
                name = "{} - {}".format(name, section_name)
            if disclaimer is not None:
                name = name + " " + disclaimer
            self.doc.add_heading(name, level=level + 1)
            # Get a binary stream for pythondocx
            imdata = io.BytesIO()
            image.save(imdata, format="png")
            imdata.seek(0)
            self._insert_img(imdata)
            # Clean the buffer
            del imdata

    def build_from_root(self, root_path: os.PathLike) -> None:
        """Given a root folder, build an atlas with all the pictures found
        in the subfolders. Only one level of subfolders is supported at the
        moment.

        Parameters
        ----------
        root_path : os.PathLike
            path to the root folder.
        """

        # TODO only one level of folders admitted for the moment
        for folder in os.listdir(root_path):
            folderpath = os.path.join(root_path, folder)

            if not os.path.isdir(folderpath):
                continue

            self.doc.add_heading(folder, level=1)

            # This ensures correct order of the pictures that will be the same
            # of creation
            files = os.listdir(folderpath)
            paths = []
            for file in files:
                paths.append(os.path.join(folderpath, file))

            for filepath in sorted(paths, key=os.path.getmtime):
                heading = os.path.basename(filepath)[:-4]
                ext = os.path.basename(filepath)[-3:]
                if ext not in ["jpg", "png", "tif"]:
                    continue

                self.doc.add_heading(heading, level=2)
                self._insert_img(filepath)

    def _change_orientation(self) -> docx.section.Section:
        current_section = self.doc.sections[-1]
        new_width, new_height = (
            current_section.page_height,
            current_section.page_width,
        )
        new_section = self.doc.add_section()
        new_section.orientation = WD_ORIENT.LANDSCAPE
        new_section.page_width = new_width
        new_section.page_height = new_height

        return new_section

    def save(
        self,
        outpath: os.PathLike,
        # pdfprint: bool = True
    ) -> None:
        """
        Save word atlas and possibly export PDF

        Parameters
        ----------
        outpath : os.PathLike
            path to the folder where to save the atlas(es)

        Returns
        -------
        None.

        """
        outpath_word = os.path.join(outpath, self.name + ".docx")
        # outpath_pdf = os.path.join(outpath, self.name+'.pdf')

        try:
            self.doc.save(outpath_word)
        except FileNotFoundError as e:
            print(" The following is the original exception:")
            print(e)
            print("\n it may be due to invalid characters in the file name")

        # if pdfprint:
        #     in_file = outpath_word
        #     out_file = outpath_pdf

        #     try:
        #         word = win32com.client.Dispatch('Word.Application')
        #     except:
        #         raise NotImplementedError('Word not installed')

        #     try:
        #         doc = word.Documents.Open(in_file)
        #         doc.ExportAsFixedFormat(
        #             OutputFileName=out_file,
        #             # 17 = PDF output, 18=XPS output
        #             ExportFormat=17,
        #             OpenAfterExport=False,
        #             # 0=Print (higher res), 1=Screen (lower res)
        #             OptimizeFor=0,
        #             # 0=No bookmarks,
        #             # 1=Heading bookmarks only,
        #             # 2=bookmarks match word bookmarks
        #             CreateBookmarks=1,
        #             DocStructureTags=True)
        #     finally:
        #         try:
        #             doc.Close()
        #         except UnboundLocalError:
        #             word.Quit()
        #             raise RuntimeError(
        #                 'something went wrong in the PDF production')
        #         word.Quit()

    # @staticmethod
    # def _wrapper(paragraph, ptype):
    #     """
    #     Wrap a paragraph in order to add cross reference

    #     Parameters
    #     ----------
    #     paragraph : docx.Paragraph
    #         image to wrap.
    #     ptype : str
    #         type of paragraph to wrap

    #     Returns
    #     -------
    #     None.

    #     """
    #     if ptype == 'table':
    #         instruction = ' SEQ Table \\* ARABIC'
    #     elif ptype == 'figure':
    #         instruction = ' SEQ Figure \\* ARABIC'
    #     else:
    #         raise ValueError(ptype+' is not a supported paragraph type')

    #     run = run = paragraph.add_run()
    #     r = run._r
    #     fldChar = OxmlElement('w:fldChar')
    #     fldChar.set(qn('w:fldCharType'), 'begin')
    #     r.append(fldChar)
    #     instrText = OxmlElement('w:instrText')
    #     instrText.text = instruction
    #     r.append(instrText)
    #     fldChar = OxmlElement('w:fldChar')
    #     fldChar.set(qn('w:fldCharType'), 'end')
    #     r.append(fldChar)

    # @staticmethod
    # def _highlightCell(cell, color='FBD4B4'):
    #     shading_elm_1 = parse_xml(r'<w:shd {} w:fill="'.format(nsdecls('w')) +
    #                               color + r'"/>')
    #     cell._tc.get_or_add_tcPr().append(shading_elm_1)


class Plotter2D(ABC):

    def __init__(
        self, suptitle: str = None, xlabel: str = None, ylabel: str = None
    ) -> None:
        """Abstract class for the definition of 2D plots.

        Parameters
        ----------
        suptitle : str, optional
            title of the plot, by default None
        xlabel : str, optional
            X axis label, by default None
        ylabel : str, optional
            Y axis label, by default None

        Attributes
        ----------
        fig : matplotlib.pyplot.Figure
            matplotlib figure representation
        ax : matplotlib.pyplot.Axis
            matplotlib axes representation
        markers : list[str]
            markers used in the plots ['o', 's', 'D', '^', 'X', 'p', 'd', '*'].
        lines : list[str]
            linestyles used in the plots ['-', '--', '-.', ':']
        colors : list[str]
            colors used in the plots. They compose a color-blind friendly
            palette.

        """
        fig, ax = plt.subplots()

        # Some default actions
        ax.grid(alpha=0.6, which="both")
        if suptitle is not None:
            ax.set_title(suptitle)
        if ylabel is not None:
            ax.set_ylabel(ylabel)
        if xlabel is not None:
            ax.set_xlabel(xlabel)

        # store the default
        self.fig = fig
        self.ax = ax

        # May be improved in the future with additional markers and colors
        # plot decorators
        self.markers = ["o", "s", "D", "^", "X", "p", "d", "*"] * 50
        self.lines = ["-", "--", "-.", ":"] * 50
        # Color-blind friendly palette
        self.colors = [
            "#377eb8",
            "#ff7f00",
            "#4daf4a",
            "#f781bf",
            "#a65628",
            "#984ea3",
            "#999999",
            "#e41a1c",
            "#dede00",
        ] * 50
        logging.debug("plotter initialized")

    @abstractmethod
    def plot(self) -> None:
        pass

    def save(self, outpath: os.PathLike) -> None:
        self.fig.savefig(outpath, dpi=300, bbox_inches="tight")
        logging.info("plot has been saved at {}".format(outpath))


class CDFplot(Plotter2D):
    def __init__(
        self, suptitle: str = None, xlabel: str = None, ylabel: str = None
    ) -> None:
        """Plotter for cumulative distributions.

        all datasets of observations are automatically binnned and plotted as
        (unfilled) histograms.

        Parameters
        ----------
        suptitle : str, optional
            title of the plot, by default None
        xlabel : str, optional
            X axis label, by default None
        ylabel : str, optional
            Y axis label, by default None

        Attributes
        ----------
        fig : matplotlib.pyplot.Figure
            matplotlib figure representation
        ax : matplotlib.pyplot.Axis
            matplotlib axes representation
        markers : list[str]
            markers used in the plots ['o', 's', 'D', '^', 'X', 'p', 'd', '*'].
        lines : list[str]
            linestyles used in the plots ['-', '--', '-.', ':']
        colors : list[str]
            colors used in the plots. They compose a color-blind friendly
            palette.

        """
        super().__init__(suptitle, xlabel, ylabel)

    def plot(
        self,
        values_list: list,
        bins: int = 10,
        datalabels: list[str] = None,
        perc: bool = True,
        outside_legend: bool = False,
        cut_y: float = None,
        cut_x: float = None,
    ) -> None:
        """plot the comulative distributions as discrete steps

        Parameters
        ----------
        values_list : list
            list of list of values containing the occurencies to be analyzed
        bins : int, optional
            number of bins into which the observations should be divided,
            by default 10
        datalabels : list[str], optional
            list of labels to be added in the plot legend representing the
            different datasets reported in values_list, by default None
        perc : bool, optional
            if True, format y-axis as percentages, by default True
        outside_legend : bool
            if True pushes the legend outside of the graph frame, by default is
            False
        cut_y : float, optional
            cut the y axis to a specific value, by default is None.
        cut_x : float, optional
            cut the x axis to a specific value, by default is None.
        """
        # check that the length of labels is consisting
        if datalabels is not None:
            try:
                assert len(datalabels) == len(values_list)
            except AssertionError:
                msg = "leghth of values ({}) is different from lenght of labels ({})"
                raise ValueError(msg.format(len(values_list), len(datalabels)))

        for i, values in enumerate(values_list):
            if datalabels is not None:
                label = datalabels[i]
            else:
                label = None

            self.ax.hist(
                values,
                bins=bins,
                histtype="step",
                weights=np.ones(len(values)) / len(values),
                cumulative=True,
                label=label,
                linestyle=self.lines[i],
                color=self.colors[i],
            )
            if perc:
                self.ax.yaxis.set_major_formatter(PercentFormatter(1))

        if outside_legend:
            self.ax.legend(framealpha=1, bbox_to_anchor=(1, 1))
        else:
            self.ax.legend(loc="lower right", framealpha=1)

        if cut_y is not None:
            self.ax.set_ylim(top=cut_y)
        if cut_x is not None:
            self.ax.set_xlim(right=cut_x)
        logging.info("CDF was plotted")
